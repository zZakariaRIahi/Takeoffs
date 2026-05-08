"""Run the full QTO pipeline locally on PDFs, photos, and DOCX files.

Usage:
    python run_local.py "24-086 Drawings.pdf"
    python run_local.py "24-086 Drawings.pdf" "24-086 Specifications.pdf"
    python run_local.py "Bid Packet.pdf" "Walkthrough Details.docx" 2874.jpg 2875.jpg
    python run_local.py "Spec Only.pdf"   # no drawings — specs path

Results saved to /tmp/local_run/results.json
"""
from __future__ import annotations

import io as _io
import json
import logging
import os
import sys
import time

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

sys.path.insert(0, os.path.dirname(__file__))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("run_local")

from app.agents.document_classifier_agent import classify_documents
from app.extractors.context_extractor import extract_context
from app.extractors.trade_extractor import extract_by_trade
from app.extractors.specs_extractor import (
    extract_from_specs_only, extract_from_text_scope, extract_from_bid_forms,
    deduplicate_items, collapse_conditionals, extract_scope_boundary, filter_excluded_items
)
from app.core.document_classification import DocumentCategory

_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
_DOCX_EXTS  = {".docx", ".doc"}
_PDF_EXTS   = {".pdf"}


def main():
    if len(sys.argv) < 2:
        print("Usage: python run_local.py <file1.pdf> [file2.pdf ...] [photo.jpg ...] [doc.docx ...]")
        sys.exit(1)

    all_files = []
    for path in sys.argv[1:]:
        if not os.path.exists(path):
            logger.error(f"File not found: {path}")
            sys.exit(1)
        with open(path, "rb") as f:
            all_files.append((os.path.basename(path), f.read()))
        logger.info(f"Loaded: {path} ({os.path.getsize(path) // 1024} KB)")

    pdf_files   = [(n, b) for n, b in all_files if os.path.splitext(n)[1].lower() in _PDF_EXTS]
    photo_files = [(n, b) for n, b in all_files if os.path.splitext(n)[1].lower() in _IMAGE_EXTS]
    docx_files  = [(n, b) for n, b in all_files if os.path.splitext(n)[1].lower() in _DOCX_EXTS]

    logger.info(f"Files: {len(pdf_files)} PDFs, {len(photo_files)} photos, {len(docx_files)} DOCX")

    # Extract DOCX text early so it can be bundled with PDF extraction
    docx_texts: list[tuple[str, str]] = []
    for fname, content in docx_files:
        try:
            from docx import Document
            doc = Document(_io.BytesIO(content))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
            if text.strip():
                docx_texts.append((fname, text))
                logger.info(f"  Parsed DOCX: {fname} ({len(text)} chars)")
        except Exception as e:
            logger.error(f"Failed to parse DOCX {fname}: {e}")

    out_dir = "/tmp/local_run"
    os.makedirs(out_dir, exist_ok=True)

    items: list = []
    scope_boundary = None

    # ── PDFs → main pipeline ──────────────────────────────────────────────────
    if pdf_files:
        logger.info("=" * 60)
        logger.info("STEP 1: Document Classification")
        logger.info("=" * 60)
        t0 = time.time()
        classification = classify_documents(pdf_files)
        logger.info(f"Step 1 done in {time.time()-t0:.1f}s")

        logger.info("=" * 60)
        logger.info("STEP 1b: Extracting scope boundary from ITB")
        logger.info("=" * 60)
        scope_boundary = extract_scope_boundary(classification)
        if scope_boundary:
            logger.info(f"Scope boundary: {len(scope_boundary.excluded)} exclusions, "
                        f"{len(scope_boundary.in_scope)} in-scope, {len(scope_boundary.alternates)} alternates")
            for ex in scope_boundary.excluded:
                logger.info(f"  EXCLUDED: {ex}")
        else:
            logger.info("No ITB files found — no scope boundary")

        has_drawings = any(
            DocumentCategory.CONSTRUCTION_DRAWINGS in (cf.categories or set())
            and (len(cf.visual_pages) / max(len(cf.pages), 1)) >= 0.10
            for cf in classification.files
        )

        if has_drawings:
            logger.info("=" * 60)
            logger.info("STEP 2: Context Extraction")
            logger.info("=" * 60)
            t0 = time.time()
            sheets, tables, rows, packages = extract_context(classification)
            logger.info(f"Step 2 done in {time.time()-t0:.1f}s — {len(sheets)} sheets, {len(tables)} tables")

            logger.info("=" * 60)
            logger.info("STEP 3: Trade Extraction")
            logger.info("=" * 60)
            t0 = time.time()
            items = extract_by_trade(classification, sheets, tables, packages)
            logger.info(f"Step 3 done in {time.time()-t0:.1f}s — {len(items)} items")

            # Also extract from specs when present — drawings alone miss spec-only scope items
            has_specs = any(
                DocumentCategory.PROJECT_SPECIFICATIONS in (cf.categories or set())
                for cf in classification.files
            )
            if has_specs or docx_texts:
                logger.info("=" * 60)
                logger.info("STEP 3b: Spec Extraction (drawings + specs path)")
                logger.info("=" * 60)
                t0 = time.time()
                spec_items = extract_from_specs_only(
                    classification, extra_text_docs=docx_texts, scope_boundary=scope_boundary
                )
                logger.info(f"Step 3b done in {time.time()-t0:.1f}s — {len(spec_items)} spec items")
                items.extend(spec_items)
            elif docx_texts:
                logger.info("DOCX: Text Scope Extraction (drawings path)")
                for fname, text in docx_texts:
                    t0 = time.time()
                    text_items = extract_from_text_scope(text, fname)
                    items.extend(text_items)
                    logger.info(f"  {fname}: {len(text_items)} items in {time.time()-t0:.1f}s")

            # Drawings path: bid forms run separately (spec extractor skips bid-form-only files)
            bid_items = extract_from_bid_forms(classification)
            if bid_items:
                items.extend(bid_items)
                logger.info(f"Bid forms: {len(bid_items)} items added")
        else:
            # Specs-only path: bundle ALL docs (PDFs + DOCX) into one call
            logger.info("=" * 60)
            logger.info("STEP 2: Specs-only Extraction (bundled)")
            logger.info("=" * 60)
            t0 = time.time()
            items = extract_from_specs_only(
                classification, extra_text_docs=docx_texts, scope_boundary=scope_boundary
            )
            logger.info(f"Specs extraction done in {time.time()-t0:.1f}s — {len(items)} items")

    # ── No PDFs: DOCX-only path ───────────────────────────────────────────────
    elif docx_texts:
        logger.info("=" * 60)
        logger.info("DOCX-ONLY: Text Scope Extraction")
        logger.info("=" * 60)
        for fname, text in docx_texts:
            t0 = time.time()
            text_items = extract_from_text_scope(text, fname)
            items.extend(text_items)
            logger.info(f"  {fname}: {len(text_items)} items in {time.time()-t0:.1f}s")

    # ── Dedup ─────────────────────────────────────────────────────────────────
    raw_path = os.path.join(out_dir, "results_raw.json")
    raw_results = [
        {
            "trade": item.trade,
            "description": item.item_description,
            "qty": item.qty,
            "unit": item.unit,
            "confidence": item.confidence,
            "source": item.source,
            "method": item.extraction_method,
            "review": item.review_reason,
            "material_spec": item.material_spec,
        }
        for item in sorted(items, key=lambda x: (x.trade, x.item_description))
    ]
    with open(raw_path, "w") as f:
        json.dump(raw_results, f, indent=2)
    logger.info(f"Raw results ({len(items)} items) saved to {raw_path}")

    logger.info("=" * 60)
    logger.info("DEDUP: Consolidating duplicates")
    logger.info("=" * 60)
    t0 = time.time()
    items = deduplicate_items(items)
    logger.info(f"Dedup done in {time.time()-t0:.1f}s")

    if scope_boundary and scope_boundary.excluded:
        logger.info("=" * 60)
        logger.info("EXCLUSION FILTER: Tagging items matching ITB exclusions")
        logger.info("=" * 60)
        items = filter_excluded_items(items, scope_boundary)

    logger.info("=" * 60)
    logger.info("COLLAPSE: Collapsing conditional / inferred items")
    logger.info("=" * 60)
    items = collapse_conditionals(items)

    # ── Save & print results ──────────────────────────────────────────────────
    results = []
    for item in sorted(items, key=lambda x: (x.trade, x.item_description)):
        results.append({
            "trade": item.trade,
            "description": item.item_description,
            "qty": item.qty,
            "unit": item.unit,
            "confidence": item.confidence,
            "source": item.source,
            "method": item.extraction_method,
            "review": item.review_reason,
            "material_spec": item.material_spec,
        })

    out_path = os.path.join(out_dir, "results.json")
    with open(out_path, "w") as f:
        json.dump(results, f, indent=2)
    logger.info(f"Results saved to {out_path}")

    print("\n" + "=" * 60)
    print(f"RESULTS — {len(results)} items")
    print("=" * 60)
    current_trade = None
    for r in results:
        if r["trade"] != current_trade:
            current_trade = r["trade"]
            print(f"\n[ {current_trade} ]")
        qty = f"{r['qty']:.0f}" if r["qty"] is not None else "TBD"
        flag = "  *** " + r["review"] if r["review"] else ""
        src = f"  [{r['source']}]" if "walkthrough_photo" in (r["source"] or "") else ""
        print(f"  {qty:>8} {r['unit']:<4}  {r['description']}{src}{flag}")

    print(f"\nFull results: {out_path}")


if __name__ == "__main__":
    main()
